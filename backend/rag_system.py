"""
RAG (Retrieval Augmented Generation) System for NPTE MCQs
Handles document processing, vector storage, and retrieval
"""

import os
import json
from typing import List, Dict, Optional, Any
from pathlib import Path
import logging

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain_community.vectorstores import Qdrant
from langchain.schema import Document
from langchain.retrievers import MultiQueryRetriever
from langchain.retrievers.document_compressors import LLMChainExtractor

# Qdrant imports
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams

# Document processing
import PyPDF2
import docx
import re

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class NPTERAGSystem:
    """RAG system specifically designed for NPTE materials"""
    
    def __init__(self, collection_name: str = "npte_materials"):
        self.collection_name = collection_name
        self.embeddings = OllamaEmbeddings(model="nomic-embed-text")
        # Use persistent storage instead of in-memory
        qdrant_path = "./qdrant_data"  # Local persistent storage
        self.qdrant_client = QdrantClient(path=qdrant_path)
        self.vector_store = None
        self.retriever = None
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,         # Optimal for research paper precision
            chunk_overlap=100,      # Minimal overlap for dense content
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
    def setup_collection(self):
        """Initialize Qdrant collection"""
        try:
            # Create collection if it doesn't exist
            collections = self.qdrant_client.get_collections()
            if self.collection_name not in [c.name for c in collections.collections]:
                self.qdrant_client.create_collection(
                    collection_name=self.collection_name,
                    vectors_config=VectorParams(size=768, distance=Distance.COSINE)  # Ollama embedding dimension
                )
                logger.info(f"Created collection: {self.collection_name}")
            else:
                logger.info(f"Collection {self.collection_name} already exists")
                
            # Initialize vector store
            self.vector_store = Qdrant(
                client=self.qdrant_client,
                collection_name=self.collection_name,
                embeddings=self.embeddings
            )
            
            # Initialize retriever (simplified for Ollama compatibility)
            self.retriever = self.vector_store.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 5}
            )
            
        except Exception as e:
            logger.error(f"Error setting up collection: {e}")
            raise
    
    def process_pdf(self, file_path: str) -> List[Document]:
        """Extract text from PDF and split into chunks"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
            # Clean text
            text = self._clean_text(text)
            
            # Split into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create documents with metadata
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "chunk_id": i,
                        "type": "pdf",
                        "topic": self._extract_topic_from_filename(file_path)
                    }
                )
                documents.append(doc)
                
            logger.info(f"Processed PDF: {file_path} -> {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return []
    
    def process_docx(self, file_path: str) -> List[Document]:
        """Extract text from DOCX and split into chunks"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Clean text
            text = self._clean_text(text)
            
            # Split into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create documents with metadata
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "chunk_id": i,
                        "type": "docx",
                        "topic": self._extract_topic_from_filename(file_path)
                    }
                )
                documents.append(doc)
                
            logger.info(f"Processed DOCX: {file_path} -> {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return []
    
    def process_text_file(self, file_path: str) -> List[Document]:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                
            # Clean text
            text = self._clean_text(text)
            
            # Split into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create documents with metadata
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "chunk_id": i,
                        "type": "txt",
                        "topic": self._extract_topic_from_filename(file_path)
                    }
                )
                documents.append(doc)
                
            logger.info(f"Processed text file: {file_path} -> {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            return []
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters that might interfere
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)]', '', text)
        return text.strip()
    
    def _extract_topic_from_filename(self, filename: str) -> str:
        """Extract topic from filename"""
        filename_lower = filename.lower()
        
        # Map filename patterns to NPTE topics
        topic_mapping = {
            "cardiovascular": "Cardiovascular and pulmonary systems",
            "cardio": "Cardiovascular and pulmonary systems",
            "musculoskeletal": "Musculoskeletal system",
            "musculo": "Musculoskeletal system",
            "neuromuscular": "Neuromuscular and nervous systems",
            "neuro": "Neuromuscular and nervous systems",
            "integumentary": "Integumentary system",
            "metabolic": "Metabolic and endocrine systems",
            "endocrine": "Metabolic and endocrine systems",
            "gastrointestinal": "Gastrointestinal system",
            "gi": "Gastrointestinal system",
            "genitourinary": "Genitourinary system",
            "gu": "Genitourinary system",
            "lymphatic": "Lymphatic system",
            "system": "System interactions"
        }
        
        for pattern, topic in topic_mapping.items():
            if pattern in filename_lower:
                return topic
                
        return "General"
    
    def add_documents(self, documents: List[Document]):
        """Add documents to the vector store"""
        try:
            if self.vector_store is None:
                self.setup_collection()
                
            # Add documents to vector store
            self.vector_store.add_documents(documents)
            logger.info(f"Added {len(documents)} documents to vector store")
            
        except Exception as e:
            logger.error(f"Error adding documents: {e}")
            raise
    
    def retrieve_relevant_context(self, query: str, topic: str = None, k: int = 5) -> List[Document]:
        """Retrieve relevant context for MCQ generation"""
        try:
            if self.vector_store is None:
                logger.warning("Vector store not initialized. Returning empty context.")
                return []
            
            # Build search query
            search_query = query
            if topic:
                search_query = f"Topic: {topic}. {query}"
            
            # Retrieve documents
            docs = self.vector_store.similarity_search(
                search_query,
                k=k,
                filter={"topic": topic} if topic else None
            )
            
            logger.info(f"Retrieved {len(docs)} relevant documents for query: {query}")
            return docs
            
        except Exception as e:
            logger.error(f"Error retrieving context: {e}")
            return []
    
    def get_context_for_mcq_generation(self, topic: str) -> str:
        """Get formatted context for MCQ generation"""
        query = f"Generate NPTE-style multiple choice questions about {topic}"
        docs = self.retrieve_relevant_context(query, topic)
        
        if not docs:
            return ""
        
        # Format context
        context_parts = []
        for i, doc in enumerate(docs):
            context_parts.append(f"Context {i+1}:\n{doc.page_content}")
        
        return "\n\n".join(context_parts)
    
    def load_documents_from_directory(self, directory_path: str):
        """Load all documents from a directory"""
        directory = Path(directory_path)
        
        if not directory.exists():
            logger.error(f"Directory does not exist: {directory_path}")
            return
        
        all_documents = []
        
        # Process all files in directory
        for file_path in directory.rglob("*"):
            if file_path.is_file():
                try:
                    if file_path.suffix.lower() == '.pdf':
                        docs = self.process_pdf(str(file_path))
                    elif file_path.suffix.lower() in ['.docx', '.doc']:
                        docs = self.process_docx(str(file_path))
                    elif file_path.suffix.lower() in ['.txt', '.md']:
                        docs = self.process_text_file(str(file_path))
                    else:
                        logger.info(f"Skipping unsupported file: {file_path}")
                        continue
                    
                    all_documents.extend(docs)
                    
                except Exception as e:
                    logger.error(f"Error processing file {file_path}: {e}")
                    continue
        
        if all_documents:
            self.add_documents(all_documents)
            logger.info(f"Successfully loaded {len(all_documents)} document chunks")
        else:
            logger.warning("No documents were successfully processed")

# Global RAG system instance
rag_system = None

def initialize_rag_system():
    """Initialize the global RAG system"""
    global rag_system
    if rag_system is None:
        rag_system = NPTERAGSystem()
        rag_system.setup_collection()
        logger.info("RAG system initialized")
    return rag_system

def get_rag_context(topic: str) -> str:
    """Get RAG context for a given topic"""
    global rag_system
    if rag_system is None:
        return ""
    
    return rag_system.get_context_for_mcq_generation(topic) 